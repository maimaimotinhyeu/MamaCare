"""
update_srs.py
Cập nhật file SRS với bảng Actors (1.3.1) và Use Cases (1.3.2)
cho hệ thống HCMS - Healthcare Clinic Management System
"""
import sys, os, copy

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ─── PATHS ────────────────────────────────────────────────────────────────────
TEMPLATE = r"e:\Documents\Ky_8\Co Dieu AISDLC\HCMS---Healthcare-Clinic-System\Template\Software Requirement Specification_v0.1.docx"
OUTPUT   = r"e:\Documents\Ky_8\Co Dieu AISDLC\HCMS---Healthcare-Clinic-System\Document\02_Requirements\HCMS_SRS_v0.1.docx"
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

# ─── COLORS (matching template: header = #FFE8E1, alt row = #FFF5F3) ──────────
HDR_BG   = "FFE8E1"   # salmon-pink header (matches template)
ALT_BG   = "FFF5F3"   # very light pink alternate row
WHITE    = "FFFFFF"
BLACK    = "000000"

# ─── DATA: ACTORS ─────────────────────────────────────────────────────────────
ACTORS = [
    ("1", "Parents\n(Phụ huynh)",
     "Human | External User",
     "Parents or guardians of pediatric patients who interact with HCMS via the self-service "
     "Booking Portal (smartphone/web). They initiate appointment bookings, provide child's "
     "Patient Demographics, receive appointment confirmations, and view post-visit clinical notes "
     "& prescriptions on their device. They are outside the clinic and do not use the internal dashboard."),

    ("2", "Doctor\n(Bác sĩ)",
     "Human | Primary Internal User & System Sponsor",
     "The lead physician (Dr. Minh) who acts as both project sponsor and the primary clinical user. "
     "Doctors interact with HCMS to access a real-time patient queue dashboard, retrieve full "
     "Medical History (EMR) for each patient, record clinical consultation outcomes into a VISIT "
     "entity, and issue electronic Prescriptions using the autocomplete drug-name feature. "
     "All Doctor interactions are within the clinic's internal network."),

    ("3", "Receptionist\n(Lễ tân)",
     "Human | Key Internal Operator",
     "The front-desk operator (Chị Lan) responsible for administrative clinic operations. "
     "Receptionists interact with HCMS to view Appointment Dashboard, confirm or manually register "
     "appointments (including walk-ins), verify the auto-generated BILLING invoices, and execute "
     "the 1-Click Payment status update (Unpaid → Paid). She is a high-frequency user of the system."),
]

# ─── DATA: USE CASES ──────────────────────────────────────────────────────────
USE_CASES = [
    # (ID, Use Case Name, Actor, Feature, Description)
    ("UC-01", "Book Appointment Online",
     "Parents",
     "Self-Service Booking Portal",
     "Parents access the online Booking Portal, select a preferred Doctor (by name/specialty) "
     "and an available time slot from the real-time calendar. The system validates the slot and "
     "auto-creates an APPOINTMENT entity with status 'Pending'. The Receptionist is notified "
     "on her dashboard of the new pending booking."),

    ("UC-02", "Provide Patient Demographics",
     "Parents",
     "Self-Service Booking Portal",
     "During an initial booking (or at check-in), Parents input the child's personal and medical "
     "data: full name, date of birth, parent phone/email, and primary medical notes. The system "
     "creates or updates the PATIENT entity linked to the APPOINTMENT record."),

    ("UC-03", "Receive Appointment Confirmation",
     "Parents",
     "Appointment Notification",
     "After the Receptionist confirms a pending appointment, HCMS automatically sends an "
     "Appointment Confirmation notification (SMS/Email) to the Parent's registered contact, "
     "including appointment date, time, doctor name, and clinic address."),

    ("UC-04", "View Clinical Notes & Prescriptions",
     "Parents",
     "Patient EMR Portal",
     "After a completed visit, Parents can log in to the portal to review the Doctor's "
     "clinical notes (VISIT record: diagnosis, symptoms, aftercare instructions) and the "
     "linked electronic PRESCRIPTION (medicine name, dosage, frequency, duration) on their smartphone."),

    ("UC-05", "View Appointment Dashboard",
     "Receptionist",
     "Appointment Management",
     "The Receptionist views a live, sortable dashboard listing all upcoming appointments for "
     "the day: patient name, time slot, assigned Doctor, and current APPOINTMENT status "
     "(Pending / Confirmed / In-Progress / Completed). The dashboard auto-refreshes in real time."),

    ("UC-06", "Confirm Appointment",
     "Receptionist",
     "Appointment Management",
     "Upon reviewing a Parent-submitted Pending appointment, the Receptionist validates the "
     "information and clicks 'Confirm'. The system updates the APPOINTMENT status from 'Pending' "
     "to 'Confirmed' and triggers the confirmation notification to the Parent (see UC-03)."),

    ("UC-07", "Register Walk-in Appointment",
     "Receptionist",
     "Appointment Management",
     "For patients who arrive at the clinic without a prior online booking, the Receptionist "
     "manually creates a new APPOINTMENT record by entering the patient info and selecting an "
     "available Doctor and time slot. The walk-in APPOINTMENT is created with status 'Confirmed'."),

    ("UC-08", "View Billing Invoice",
     "Receptionist",
     "Billing & Payment",
     "Once the Doctor completes and saves a VISIT + PRESCRIPTION record, the HCMS system "
     "automatically generates a BILLING invoice summarizing: consultation fee, itemized "
     "prescription costs, and total amount due. The Receptionist can view this invoice on "
     "the dashboard with status 'Unpaid'."),

    ("UC-09", "Update Payment Status (1-Click Paid)",
     "Receptionist",
     "Billing & Payment",
     "After collecting the payment from the Parent (cash or bank transfer at the counter), the "
     "Receptionist clicks a single 'Mark as Paid' button. The system instantly updates the "
     "BILLING status from 'Unpaid' to 'Paid', records the payment timestamp, and logs the "
     "transaction for the end-of-day revenue reconciliation report."),

    ("UC-10", "View Patient Queue",
     "Doctor",
     "Clinical Dashboard",
     "The Doctor opens their personal clinical dashboard to see the ordered list of patients "
     "waiting for examination: patient name, appointment time, and current status. The queue "
     "is auto-sorted by appointment time and updated in real time as Receptionists confirm arrivals."),

    ("UC-11", "Access Patient Medical History (EMR)",
     "Doctor",
     "EMR – Medical History",
     "Before or during an examination, the Doctor selects a patient from the queue and retrieves "
     "their complete Electronic Medical Record: full PATIENT profile, all previous VISIT records "
     "(past diagnoses, symptoms), and all past PRESCRIPTION records. This enables evidence-based, "
     "data-driven clinical decisions without paper records."),

    ("UC-12", "Record Clinical Consultation (VISIT)",
     "Doctor",
     "EMR – Visit Management",
     "During the physical examination, the Doctor enters the consultation outcome into a new VISIT "
     "record: observed symptoms, clinical diagnosis (ICD-coded), physical examination notes, and "
     "aftercare/follow-up instructions. The VISIT record is saved and linked to the patient's "
     "PATIENT entity and the originating APPOINTMENT."),

    ("UC-13", "Issue Electronic Prescription",
     "Doctor",
     "EMR – Prescription Management",
     "After completing the VISIT record, the Doctor creates a PRESCRIPTION record by specifying "
     "each medicine (drug name with autocomplete suggestion), dosage amount, frequency, "
     "administration route, and duration. The saved PRESCRIPTION is automatically linked to "
     "the VISIT and PATIENT entities, and is visible to the Parent via the portal (UC-04)."),

    ("UC-14", "Manage Patient Profile",
     "Doctor",
     "EMR – Patient Management",
     "The Doctor can view and update supplementary medical information in the PATIENT profile: "
     "known allergies, blood type, chronic condition flags, and vaccination notes. This ensures "
     "the EMR remains accurate and up-to-date for future clinical decisions by the same or "
     "other doctors at the clinic."),
]

# ─── HELPER FUNCTIONS ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Set background shading of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for old_shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(old_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_padding(cell, top=50, bottom=50, left=108, right=108):
    """Set cell internal padding (in twentieths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcPr.append(tcMar)

def add_cell_text(cell, text, bold=False, italic=False, font_size_pt=None,
                  color_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear cell and add a formatted run."""
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)

def set_col_width(table, col_idx, width_cm):
    """Set column width via XML (fallback for autofit columns)."""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def set_table_no_autofit(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    # Remove existing tblLayout
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblPr.append(tblLayout)

def build_actors_table(doc, para_index):
    """Insert the Actors table after the paragraph at para_index."""
    # Build table: 4 rows (header + 3 actors), 3 cols (#, Actor, Description)
    table = doc.add_table(rows=1 + len(ACTORS), cols=3)
    table.style = 'Table Grid'
    set_table_no_autofit(table)

    # Column widths: # (1cm), Actor (4cm), Description (11cm)
    widths = [1.0, 4.5, 10.5]
    for row in table.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = Cm(w)

    # ── HEADER ROW ──
    headers = ["#", "Actor", "Description"]
    hrow = table.rows[0]
    for ci, hdr in enumerate(headers):
        cell = hrow.cells[ci]
        set_cell_bg(cell, HDR_BG)
        set_cell_padding(cell)
        add_cell_text(cell, hdr, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── DATA ROWS ──
    for ri, (no, actor, atype, desc) in enumerate(ACTORS):
        row = table.rows[ri + 1]
        bg = WHITE if ri % 2 == 0 else ALT_BG

        # Col 0: #
        set_cell_bg(row.cells[0], bg)
        set_cell_padding(row.cells[0])
        add_cell_text(row.cells[0], no, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Col 1: Actor Name + Type
        set_cell_bg(row.cells[1], bg)
        set_cell_padding(row.cells[1])
        cell = row.cells[1]
        cell.text = ''
        para = cell.paragraphs[0]
        r1 = para.add_run(actor)
        r1.bold = True
        para.add_run('\n')
        r2 = para.add_run(atype)
        r2.italic = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor.from_string("555555")

        # Col 2: Description
        set_cell_bg(row.cells[2], bg)
        set_cell_padding(row.cells[2])
        add_cell_text(row.cells[2], desc)

    # ── Move table to correct position via XML ──
    # Get the paragraph XML element to insert after
    target_para = doc.paragraphs[para_index]._element
    target_para.addnext(table._tbl)
    print(f"  ✅ Actors table inserted after paragraph [{para_index}]")
    return table

def build_usecases_table(doc, para_index):
    """Insert the Use Cases table after the paragraph at para_index."""
    # Table: header + len(USE_CASES) rows, 4 cols
    table = doc.add_table(rows=1 + len(USE_CASES), cols=4)
    table.style = 'Table Grid'
    set_table_no_autofit(table)

    # Column widths: ID(1.5), Use Case(4), Feature(3.5), Description(7)
    widths = [1.5, 4.0, 3.5, 7.0]
    for row in table.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = Cm(w)

    # ── HEADER ROW ──
    headers = ["ID", "Use Case", "Feature", "Use Case Description"]
    hrow = table.rows[0]
    for ci, hdr in enumerate(headers):
        cell = hrow.cells[ci]
        set_cell_bg(cell, HDR_BG)
        set_cell_padding(cell)
        add_cell_text(cell, hdr, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── DATA ROWS ──
    for ri, (uc_id, uc_name, actor, feature, desc) in enumerate(USE_CASES):
        row = table.rows[ri + 1]
        # Group rows by actor for visual separation
        if actor == "Parents":
            bg = WHITE if ri % 2 == 0 else ALT_BG
        elif actor == "Receptionist":
            bg = "FFF0F5" if ri % 2 == 0 else "FFEBEE"
        else:  # Doctor
            bg = "F3FFF3" if ri % 2 == 0 else "EBFFEB"

        # Col 0: ID
        set_cell_bg(row.cells[0], bg)
        set_cell_padding(row.cells[0])
        add_cell_text(row.cells[0], uc_id, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Col 1: Use Case Name
        set_cell_bg(row.cells[1], bg)
        set_cell_padding(row.cells[1])
        add_cell_text(row.cells[1], uc_name, bold=True)

        # Col 2: Feature
        set_cell_bg(row.cells[2], bg)
        set_cell_padding(row.cells[2])
        add_cell_text(row.cells[2], feature, italic=True)

        # Col 3: Description
        set_cell_bg(row.cells[3], bg)
        set_cell_padding(row.cells[3])
        add_cell_text(row.cells[3], desc)

    # Move table to correct position
    target_para = doc.paragraphs[para_index]._element
    target_para.addnext(table._tbl)
    print(f"  ✅ Use Cases table inserted after paragraph [{para_index}]")
    return table

# ─── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    print("[*] Opening SRS template...")
    doc = Document(TEMPLATE)

    # ── Find section 1.3.1 Actors ──
    actors_para_idx  = None
    uc_para_idx      = None

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if '1.3.1' in text and 'Actor' in text:
            actors_para_idx = i
        if '1.3.2' in text and 'Use Case' in text and 'Diagram' not in text:
            if uc_para_idx is None:
                uc_para_idx = i

    print(f"  → 1.3.1 Actors heading at paragraph [{actors_para_idx}]")
    print(f"  → 1.3.2 Use Cases heading at paragraph [{uc_para_idx}]")

    if actors_para_idx is None or uc_para_idx is None:
        print("❌ Could not find section headings. Aborting.")
        return

    # Remove existing template Actor/UC tables (Table[1] and Table[2])
    # We'll re-insert our own styled tables after the headings
    # Find the last guidance paragraph in 1.3.1 section (before 1.3.2 heading)
    # Guidance paragraphs are paras [40]-[50] in the template
    # We insert after the last guidance paragraph (index 49)
    actors_insert_after = None
    for i in range(actors_para_idx, uc_para_idx):
        t = doc.paragraphs[i].text.strip()
        if 'table form as below' in t.lower():
            actors_insert_after = i
            break
    if actors_insert_after is None:
        atoms_insert_after = actors_para_idx

    # For use cases, insert after the "table form as below" paragraph in 1.3.2 section
    uc_insert_after = None
    for i in range(uc_para_idx, min(uc_para_idx + 10, len(doc.paragraphs))):
        t = doc.paragraphs[i].text.strip()
        if 'table form as below' in t.lower() or '•' in t or 'below]' in t.lower():
            uc_insert_after = i
    if uc_insert_after is None:
        uc_insert_after = uc_para_idx

    print(f"\n🔵 Building Actors table (inserting after para [{actors_insert_after}])...")
    actors_tbl = build_actors_table(doc, actors_insert_after)

    # After inserting actors table, paragraph indices shift - recalculate uc_insert_after
    # (The table was inserted into the body XML directly, paragraph count stays the same)
    print(f"\n🔵 Building Use Cases table (inserting after para [{uc_insert_after}])...")
    uc_tbl = build_usecases_table(doc, uc_insert_after)

    # Save
    doc.save(OUTPUT)
    print(f"\n✅ SRS document saved to:\n   {OUTPUT}")

if __name__ == '__main__':
    main()
